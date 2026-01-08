"""
Resume Generator Module
Generates updated resume with new points added in appropriate sections
"""

import os
from typing import Dict, List
from datetime import datetime


class ResumeGenerator:
    """Generate updated resume with missing points added"""
    
    def __init__(self):
        pass
    
    def generate(self, resume_data: Dict, added_points: List[Dict], 
                 original_filepath: str, output_filepath: str) -> str:
        """
        Generate an updated resume with added points
        
        Args:
            resume_data: Original resume data
            added_points: List of points to add with their sections
            original_filepath: Path to original resume file
            output_filepath: Path where to save updated resume
            
        Returns:
            Path to generated resume
        """
        # Organize points by section
        points_by_section = self._organize_points_by_section(added_points)
        
        # Get original file extension
        file_extension = original_filepath.rsplit('.', 1)[1].lower()
        
        # Generate based on file type
        if file_extension == 'pdf':
            return self._generate_pdf(resume_data, points_by_section, output_filepath)
        elif file_extension == 'docx':
            return self._generate_docx(resume_data, points_by_section, output_filepath)
        else:  # txt or fallback
            return self._generate_txt(resume_data, points_by_section, output_filepath)
    
    def _organize_points_by_section(self, added_points: List[Dict]) -> Dict[str, List[Dict]]:
        """Organize added points by their target section"""
        organized = {}
        
        for point in added_points:
            section = point.get('section', 'Additional Information')
            if section not in organized:
                organized[section] = []
            organized[section].append(point)
        
        return organized
    
    def _generate_pdf(self, resume_data: Dict, points_by_section: Dict, 
                      output_filepath: str) -> str:
        """Generate PDF resume"""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib.enums import TA_LEFT, TA_CENTER
            
            # Create PDF document
            doc = SimpleDocTemplate(output_filepath, pagesize=letter,
                                   rightMargin=0.75*inch, leftMargin=0.75*inch,
                                   topMargin=0.75*inch, bottomMargin=0.75*inch)
            
            # Container for the 'Flowable' objects
            elements = []
            
            # Define styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                textColor='#000000',
                spaceAfter=12,
                alignment=TA_CENTER
            )
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                textColor='#000000',
                spaceAfter=6,
                spaceBefore=12
            )
            normal_style = styles['Normal']
            bullet_style = ParagraphStyle(
                'Bullet',
                parent=styles['Normal'],
                leftIndent=20,
                spaceAfter=6
            )
            
            # Add title
            elements.append(Paragraph("Updated Resume", title_style))
            elements.append(Spacer(1, 0.2*inch))
            
            # Process each section
            sections = resume_data.get('sections', [])
            
            for section in sections:
                section_name = section.get('name', 'Section')
                section_content = section.get('content', '')
                
                # Add section heading
                elements.append(Paragraph(section_name.upper(), heading_style))
                
                # Add original content
                if section_content:
                    content_lines = section_content.split('\n')
                    for line in content_lines:
                        if line.strip():
                            elements.append(Paragraph(f"• {line.strip()}", bullet_style))
                
                # Add new points for this section
                section_points = points_by_section.get(section_name, [])
                for point in section_points:
                    point_text = point.get('point', '')
                    project = point.get('project', '')
                    additional_info = point.get('additional_info', '')
                    
                    # Format the point
                    formatted_point = self._format_point_for_display(
                        point_text, project, additional_info
                    )
                    elements.append(Paragraph(f"• {formatted_point}", bullet_style))
                
                elements.append(Spacer(1, 0.1*inch))
            
            # Add any points that didn't match existing sections
            unmatched_sections = set(points_by_section.keys()) - set([s['name'] for s in sections])
            for section_name in unmatched_sections:
                elements.append(Paragraph(section_name.upper(), heading_style))
                
                for point in points_by_section[section_name]:
                    point_text = point.get('point', '')
                    project = point.get('project', '')
                    additional_info = point.get('additional_info', '')
                    
                    formatted_point = self._format_point_for_display(
                        point_text, project, additional_info
                    )
                    elements.append(Paragraph(f"• {formatted_point}", bullet_style))
                
                elements.append(Spacer(1, 0.1*inch))
            
            # Build PDF
            doc.build(elements)
            
            return output_filepath
            
        except ImportError:
            # Fallback to text if reportlab is not available
            return self._generate_txt(resume_data, points_by_section, output_filepath)
    
    def _generate_docx(self, resume_data: Dict, points_by_section: Dict, 
                       output_filepath: str) -> str:
        """Generate DOCX resume"""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            
            doc = Document()
            
            # Add title
            title = doc.add_heading('Updated Resume', 0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Process each section
            sections = resume_data.get('sections', [])
            
            for section in sections:
                section_name = section.get('name', 'Section')
                section_content = section.get('content', '')
                
                # Add section heading
                doc.add_heading(section_name.upper(), level=1)
                
                # Add original content
                if section_content:
                    content_lines = section_content.split('\n')
                    for line in content_lines:
                        if line.strip():
                            p = doc.add_paragraph(line.strip(), style='List Bullet')
                
                # Add new points for this section
                section_points = points_by_section.get(section_name, [])
                for point in section_points:
                    point_text = point.get('point', '')
                    project = point.get('project', '')
                    additional_info = point.get('additional_info', '')
                    
                    formatted_point = self._format_point_for_display(
                        point_text, project, additional_info
                    )
                    p = doc.add_paragraph(formatted_point, style='List Bullet')
                    # Highlight new points
                    run = p.runs[0]
                    run.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue for new points
            
            # Add any points that didn't match existing sections
            unmatched_sections = set(points_by_section.keys()) - set([s['name'] for s in sections])
            for section_name in unmatched_sections:
                doc.add_heading(section_name.upper(), level=1)
                
                for point in points_by_section[section_name]:
                    point_text = point.get('point', '')
                    project = point.get('project', '')
                    additional_info = point.get('additional_info', '')
                    
                    formatted_point = self._format_point_for_display(
                        point_text, project, additional_info
                    )
                    p = doc.add_paragraph(formatted_point, style='List Bullet')
                    run = p.runs[0]
                    run.font.color.rgb = RGBColor(0, 0, 139)
            
            # Save document
            doc.save(output_filepath)
            
            return output_filepath
            
        except ImportError:
            # Fallback to text if python-docx is not available
            return self._generate_txt(resume_data, points_by_section, output_filepath)
    
    def _generate_txt(self, resume_data: Dict, points_by_section: Dict, 
                      output_filepath: str) -> str:
        """Generate TXT resume"""
        output_lines = []
        
        output_lines.append("=" * 80)
        output_lines.append("UPDATED RESUME")
        output_lines.append(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        output_lines.append("=" * 80)
        output_lines.append("")
        
        # Process each section
        sections = resume_data.get('sections', [])
        
        for section in sections:
            section_name = section.get('name', 'Section')
            section_content = section.get('content', '')
            
            # Add section heading
            output_lines.append(section_name.upper())
            output_lines.append("-" * len(section_name))
            
            # Add original content
            if section_content:
                content_lines = section_content.split('\n')
                for line in content_lines:
                    if line.strip():
                        output_lines.append(f"• {line.strip()}")
            
            # Add new points for this section
            section_points = points_by_section.get(section_name, [])
            for point in section_points:
                point_text = point.get('point', '')
                project = point.get('project', '')
                additional_info = point.get('additional_info', '')
                
                formatted_point = self._format_point_for_display(
                    point_text, project, additional_info
                )
                output_lines.append(f"• [NEW] {formatted_point}")
            
            output_lines.append("")
        
        # Add any points that didn't match existing sections
        unmatched_sections = set(points_by_section.keys()) - set([s['name'] for s in sections])
        for section_name in unmatched_sections:
            output_lines.append(section_name.upper())
            output_lines.append("-" * len(section_name))
            
            for point in points_by_section[section_name]:
                point_text = point.get('point', '')
                project = point.get('project', '')
                additional_info = point.get('additional_info', '')
                
                formatted_point = self._format_point_for_display(
                    point_text, project, additional_info
                )
                output_lines.append(f"• [NEW] {formatted_point}")
            
            output_lines.append("")
        
        # Write to file
        with open(output_filepath, 'w', encoding='utf-8') as f:
            f.write('\n'.join(output_lines))
        
        return output_filepath
    
    def _format_point_for_display(self, point: str, project: str, additional_info: str) -> str:
        """Format a point for display in the resume"""
        formatted = point
        
        if project:
            formatted = f"{formatted} (Project: {project})"
        
        if additional_info:
            formatted = f"{formatted} - {additional_info}"
        
        return formatted
