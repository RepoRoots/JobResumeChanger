"""
Resume Parser Module
Extracts text and sections from PDF and DOCX resume files.
"""

import re
from pathlib import Path

import fitz  # PyMuPDF for PDF parsing
from docx import Document


class ResumeParser:
    """Parse resume files (PDF/DOCX) and extract structured content."""

    # Common section headers in resumes
    SECTION_PATTERNS = [
        r'(?i)^(professional\s+)?summary',
        r'(?i)^(career\s+)?objective',
        r'(?i)^experience',
        r'(?i)^work\s+(experience|history)',
        r'(?i)^professional\s+experience',
        r'(?i)^employment(\s+history)?',
        r'(?i)^education',
        r'(?i)^(academic\s+)?qualifications',
        r'(?i)^skills',
        r'(?i)^(technical\s+)?skills',
        r'(?i)^core\s+competencies',
        r'(?i)^projects',
        r'(?i)^(personal\s+)?projects',
        r'(?i)^certifications?',
        r'(?i)^licenses?\s*(and|&)?\s*certifications?',
        r'(?i)^achievements?',
        r'(?i)^accomplishments?',
        r'(?i)^awards?(\s*(and|&)\s*honors?)?',
        r'(?i)^publications?',
        r'(?i)^languages?',
        r'(?i)^interests?',
        r'(?i)^hobbies',
        r'(?i)^references?',
        r'(?i)^volunteer',
        r'(?i)^leadership',
        r'(?i)^activities',
        r'(?i)^extracurricular',
    ]

    def __init__(self):
        self.section_regex = re.compile(
            '|'.join(f'({p})' for p in self.SECTION_PATTERNS)
        )

    def parse(self, filepath: str) -> tuple[str, dict]:
        """
        Parse a resume file and extract text and sections.

        Args:
            filepath: Path to the resume file (PDF or DOCX)

        Returns:
            Tuple of (full_text, sections_dict)
        """
        filepath = Path(filepath)
        ext = filepath.suffix.lower()

        if ext == '.pdf':
            text = self._parse_pdf(filepath)
        elif ext == '.docx':
            text = self._parse_docx(filepath)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

        sections = self._extract_sections(text)
        return text, sections

    def _parse_pdf(self, filepath: Path) -> str:
        """Extract text from PDF file."""
        text_parts = []

        with fitz.open(filepath) as doc:
            for page in doc:
                text_parts.append(page.get_text())

        return '\n'.join(text_parts)

    def _parse_docx(self, filepath: Path) -> str:
        """Extract text from DOCX file."""
        doc = Document(filepath)
        text_parts = []

        for para in doc.paragraphs:
            text_parts.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)

        return '\n'.join(text_parts)

    def _extract_sections(self, text: str) -> dict:
        """
        Extract sections from resume text.

        Returns:
            Dictionary mapping section names to their content
        """
        lines = text.split('\n')
        sections = {}
        current_section = 'header'
        current_content = []

        for line in lines:
            stripped = line.strip()

            # Check if this line is a section header
            if stripped and self._is_section_header(stripped):
                # Save previous section
                if current_content:
                    sections[current_section] = '\n'.join(current_content).strip()

                # Start new section
                current_section = self._normalize_section_name(stripped)
                current_content = []
            else:
                current_content.append(line)

        # Save the last section
        if current_content:
            sections[current_section] = '\n'.join(current_content).strip()

        return sections

    def _is_section_header(self, text: str) -> bool:
        """Check if a line is a section header."""
        # Remove common formatting characters
        cleaned = re.sub(r'[:\-_|]', '', text).strip()

        # Check if it matches section patterns
        if self.section_regex.match(cleaned):
            return True

        # Additional heuristics for section headers
        # - Short lines (typically < 50 chars)
        # - All caps or title case
        # - Ends with colon
        if len(cleaned) < 50:
            if cleaned.isupper() or cleaned.istitle():
                if self.section_regex.search(cleaned):
                    return True

        return False

    def _normalize_section_name(self, header: str) -> str:
        """Normalize section header to standard name."""
        header_lower = header.lower().strip()

        # Remove common punctuation
        header_clean = re.sub(r'[:\-_|]', '', header_lower).strip()

        # Map to standard section names
        section_mapping = {
            'summary': ['summary', 'professional summary', 'career summary', 'profile'],
            'objective': ['objective', 'career objective'],
            'experience': ['experience', 'work experience', 'work history',
                          'professional experience', 'employment', 'employment history'],
            'education': ['education', 'academic qualifications', 'qualifications'],
            'skills': ['skills', 'technical skills', 'core competencies', 'competencies'],
            'projects': ['projects', 'personal projects', 'academic projects'],
            'certifications': ['certifications', 'certification', 'licenses',
                              'licenses and certifications', 'licenses & certifications'],
            'achievements': ['achievements', 'accomplishments', 'awards',
                            'awards and honors', 'honors'],
            'publications': ['publications', 'research', 'papers'],
            'languages': ['languages', 'language skills'],
            'interests': ['interests', 'hobbies', 'personal interests'],
            'references': ['references'],
            'volunteer': ['volunteer', 'volunteer experience', 'volunteering'],
            'leadership': ['leadership', 'leadership experience'],
            'activities': ['activities', 'extracurricular', 'extracurricular activities'],
        }

        for standard_name, variations in section_mapping.items():
            if header_clean in variations:
                return standard_name

        # If no match, return cleaned header
        return header_clean


def extract_bullet_points(text: str) -> list[str]:
    """Extract bullet points from text."""
    bullet_patterns = [
        r'^[\u2022\u2023\u25E6\u2043\u2219]\s*',  # Unicode bullets
        r'^[-*+]\s+',  # Common ASCII bullets
        r'^\d+[.)]\s+',  # Numbered lists
        r'^[a-zA-Z][.)]\s+',  # Lettered lists
    ]

    combined_pattern = '|'.join(f'({p})' for p in bullet_patterns)
    bullet_regex = re.compile(combined_pattern, re.MULTILINE)

    points = []
    lines = text.split('\n')

    for line in lines:
        stripped = line.strip()
        if stripped and bullet_regex.match(stripped):
            # Remove the bullet character
            clean = bullet_regex.sub('', stripped).strip()
            if clean:
                points.append(clean)

    return points
