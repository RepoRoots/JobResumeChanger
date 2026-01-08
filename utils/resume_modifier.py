"""
Resume Modifier Module
Modifies resume files (PDF/DOCX) with new content.
"""

import os
import re
import copy
from pathlib import Path
from typing import Optional
from dataclasses import dataclass

import fitz  # PyMuPDF for PDF
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


@dataclass
class Modification:
    """A single modification to apply to the resume."""
    section: str
    content: str
    position: str = 'end'  # 'start', 'end', or 'after:text'
    style: Optional[str] = None


class ResumeModifier:
    """Modify resume files with new content."""

    def __init__(self):
        self.bullet_char = '•'

    def generate_content(
        self,
        missing_point: str,
        search_info: str,
        target_section: str,
        project_context: str = ''
    ) -> str:
        """
        Generate resume content for a missing point.

        Args:
            missing_point: The skill/requirement to add
            search_info: Information gathered from web search
            target_section: Which section to add content to
            project_context: Optional project context provided by user

        Returns:
            Generated content string
        """
        # Determine content type based on section
        if target_section.lower() in ['skills', 'technical skills', 'core competencies']:
            # For skills section, just return the keyword
            return missing_point

        elif target_section.lower() in ['experience', 'work experience', 'professional experience']:
            # Generate a bullet point for experience
            return self._generate_experience_bullet(missing_point, project_context, search_info)

        elif target_section.lower() in ['projects', 'personal projects']:
            # Generate project description
            return self._generate_project_bullet(missing_point, project_context, search_info)

        elif target_section.lower() in ['summary', 'professional summary', 'objective']:
            # Generate summary enhancement
            return self._generate_summary_addition(missing_point, search_info)

        else:
            # Generic bullet point
            return self._generate_generic_bullet(missing_point, project_context)

    def _generate_experience_bullet(
        self,
        skill: str,
        project_context: str,
        search_info: str
    ) -> str:
        """Generate a bullet point for experience section."""
        action_verbs = [
            'Developed', 'Implemented', 'Designed', 'Built', 'Created',
            'Engineered', 'Architected', 'Optimized', 'Enhanced', 'Delivered'
        ]

        if project_context:
            return f"{action_verbs[0]} {skill} solutions for {project_context}, improving efficiency and delivering measurable results"
        return f"{action_verbs[0]} and maintained {skill}-based applications, ensuring high performance and reliability"

    def _generate_project_bullet(
        self,
        skill: str,
        project_context: str,
        search_info: str
    ) -> str:
        """Generate content for projects section."""
        if project_context:
            return f"Utilized {skill} to build {project_context}, demonstrating practical expertise and problem-solving abilities"
        return f"Applied {skill} in personal project to explore advanced concepts and best practices"

    def _generate_summary_addition(self, skill: str, search_info: str) -> str:
        """Generate addition for summary section."""
        return f"Proficient in {skill}"

    def _generate_generic_bullet(self, skill: str, context: str) -> str:
        """Generate a generic bullet point."""
        if context:
            return f"Applied {skill} expertise in {context}"
        return f"Experienced with {skill}"

    def apply_modifications(
        self,
        input_path: str,
        output_path: str,
        modifications: list[dict],
        file_ext: str
    ) -> None:
        """
        Apply modifications to resume and save to output path.

        Args:
            input_path: Path to original resume
            output_path: Path for modified resume
            modifications: List of modifications to apply
            file_ext: File extension (pdf or docx)
        """
        if file_ext.lower() == 'docx':
            self._modify_docx(input_path, output_path, modifications)
        elif file_ext.lower() == 'pdf':
            self._modify_pdf(input_path, output_path, modifications)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")

    def _modify_docx(
        self,
        input_path: str,
        output_path: str,
        modifications: list[dict]
    ) -> None:
        """Modify a DOCX resume."""
        doc = Document(input_path)

        # Group modifications by section
        section_mods = {}
        for mod in modifications:
            section = mod.get('section', '').lower()
            if section not in section_mods:
                section_mods[section] = []
            section_mods[section].append(mod)

        # Find and modify sections
        current_section = None
        section_paragraphs = {}

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip().lower()

            # Check if this is a section header
            if self._is_section_header(text):
                current_section = self._normalize_section(text)
                section_paragraphs[current_section] = {'start': i, 'end': i}
            elif current_section:
                section_paragraphs[current_section]['end'] = i

        # Apply modifications
        paragraphs_to_add = []

        for section, mods in section_mods.items():
            section_key = self._normalize_section(section)

            if section_key in section_paragraphs:
                insert_index = section_paragraphs[section_key]['end'] + 1

                for mod in mods:
                    content = mod.get('content', '')
                    if content:
                        paragraphs_to_add.append({
                            'index': insert_index,
                            'content': content,
                            'is_bullet': section_key not in ['skills', 'summary']
                        })
                        insert_index += 1
            else:
                # Section not found, add to end of document
                for mod in mods:
                    content = mod.get('content', '')
                    if content:
                        paragraphs_to_add.append({
                            'index': len(doc.paragraphs),
                            'content': content,
                            'is_bullet': True
                        })

        # Sort by index in reverse order to maintain correct positions
        paragraphs_to_add.sort(key=lambda x: x['index'], reverse=True)

        # Create a new document with modifications
        new_doc = Document(input_path)

        # Track which paragraphs have been processed
        for add_info in paragraphs_to_add:
            idx = min(add_info['index'], len(new_doc.paragraphs))

            # Add new paragraph
            if idx < len(new_doc.paragraphs):
                # Insert after the target paragraph
                target_para = new_doc.paragraphs[idx - 1] if idx > 0 else new_doc.paragraphs[0]
                new_para = self._insert_paragraph_after(target_para, add_info['content'])

                if add_info['is_bullet']:
                    # Format as bullet point
                    new_para.text = f"{self.bullet_char} {add_info['content']}"
            else:
                # Append to end
                new_para = new_doc.add_paragraph()
                if add_info['is_bullet']:
                    new_para.text = f"{self.bullet_char} {add_info['content']}"
                else:
                    new_para.text = add_info['content']

        new_doc.save(output_path)

    def _insert_paragraph_after(self, paragraph, text: str):
        """Insert a new paragraph after an existing one."""
        new_para = paragraph.insert_paragraph_before(text)
        # Move the new paragraph after
        paragraph._p.addnext(new_para._p)
        return new_para

    def _modify_pdf(
        self,
        input_path: str,
        output_path: str,
        modifications: list[dict]
    ) -> None:
        """
        Modify a PDF resume.

        Note: PDF modification is limited. We create annotations/overlays
        or convert to a new format. For best results, DOCX is preferred.
        """
        doc = fitz.open(input_path)

        # Group modifications by section
        section_mods = {}
        for mod in modifications:
            section = mod.get('section', '').lower()
            if section not in section_mods:
                section_mods[section] = []
            section_mods[section].append(mod)

        # Find section positions in PDF
        section_positions = self._find_pdf_sections(doc)

        # Apply modifications
        for section, mods in section_mods.items():
            section_key = self._normalize_section(section)

            if section_key in section_positions:
                pos_info = section_positions[section_key]
                page_num = pos_info['page']
                y_pos = pos_info['y_end']

                page = doc[page_num]

                for mod in mods:
                    content = mod.get('content', '')
                    if content:
                        # Add text annotation or overlay
                        y_pos += 15  # Move down for each new line

                        # Format as bullet if not skills/summary
                        if section_key not in ['skills', 'summary']:
                            text = f"• {content}"
                        else:
                            text = content

                        # Insert text
                        text_point = fitz.Point(72, y_pos)  # 72 points = 1 inch margin

                        # Check if we need a new page
                        if y_pos > page.rect.height - 72:
                            page = doc.new_page(-1, width=page.rect.width, height=page.rect.height)
                            y_pos = 72

                        page.insert_text(
                            text_point,
                            text,
                            fontsize=10,
                            fontname="helv"
                        )
            else:
                # Add to last page
                last_page = doc[-1]
                y_pos = last_page.rect.height - 100

                for mod in mods:
                    content = mod.get('content', '')
                    if content:
                        text = f"• {content}"
                        text_point = fitz.Point(72, y_pos)
                        last_page.insert_text(
                            text_point,
                            text,
                            fontsize=10,
                            fontname="helv"
                        )
                        y_pos += 15

        doc.save(output_path)
        doc.close()

    def _find_pdf_sections(self, doc) -> dict:
        """Find section positions in PDF document."""
        sections = {}

        section_keywords = [
            'summary', 'objective', 'experience', 'education',
            'skills', 'projects', 'certifications', 'achievements'
        ]

        for page_num, page in enumerate(doc):
            blocks = page.get_text("dict")["blocks"]

            for block in blocks:
                if "lines" in block:
                    for line in block["lines"]:
                        for span in line["spans"]:
                            text = span["text"].strip().lower()

                            for keyword in section_keywords:
                                if keyword in text and len(text) < 50:
                                    section_key = self._normalize_section(text)
                                    bbox = span["bbox"]

                                    if section_key not in sections:
                                        sections[section_key] = {
                                            'page': page_num,
                                            'y_start': bbox[1],
                                            'y_end': bbox[3]
                                        }

        # Update y_end for each section based on next section
        section_list = sorted(sections.items(), key=lambda x: (x[1]['page'], x[1]['y_start']))

        for i, (name, pos) in enumerate(section_list):
            if i + 1 < len(section_list):
                next_pos = section_list[i + 1][1]
                if next_pos['page'] == pos['page']:
                    sections[name]['y_end'] = next_pos['y_start'] - 10

        return sections

    def _is_section_header(self, text: str) -> bool:
        """Check if text is a section header."""
        section_keywords = [
            'summary', 'objective', 'experience', 'education',
            'skills', 'projects', 'certifications', 'achievements',
            'qualifications', 'work history', 'employment', 'technical'
        ]

        text_lower = text.lower().strip()

        # Short text with section keyword
        if len(text_lower) < 50:
            for keyword in section_keywords:
                if keyword in text_lower:
                    return True

        return False

    def _normalize_section(self, text: str) -> str:
        """Normalize section name."""
        text_lower = text.lower().strip()

        # Remove common punctuation
        text_clean = re.sub(r'[:\-_|]', '', text_lower).strip()

        # Map to standard names
        mapping = {
            'summary': ['summary', 'professional summary', 'career summary', 'profile'],
            'experience': ['experience', 'work experience', 'work history', 'employment',
                          'professional experience'],
            'education': ['education', 'academic', 'qualifications'],
            'skills': ['skills', 'technical skills', 'core competencies', 'technologies'],
            'projects': ['projects', 'personal projects', 'key projects'],
            'certifications': ['certifications', 'certificates', 'licenses'],
            'achievements': ['achievements', 'accomplishments', 'awards', 'honors'],
        }

        for standard, variations in mapping.items():
            if any(var in text_clean for var in variations):
                return standard

        return text_clean


class ContentFormatter:
    """Format content for resume insertion."""

    @staticmethod
    def format_skill_list(skills: list[str], separator: str = ' | ') -> str:
        """Format skills as a separated list."""
        return separator.join(skills)

    @staticmethod
    def format_bullet_point(content: str, verb: Optional[str] = None) -> str:
        """Format content as a resume bullet point."""
        if verb and not content.startswith(verb):
            # Capitalize first letter after verb
            content = f"{verb} {content[0].lower()}{content[1:]}"

        # Ensure it starts with capital letter
        content = content[0].upper() + content[1:]

        # Remove trailing period if present, add one
        content = content.rstrip('.') + '.'

        return content

    @staticmethod
    def format_for_section(content: str, section: str) -> str:
        """Format content appropriately for a specific section."""
        section_lower = section.lower()

        if 'skill' in section_lower:
            # Skills are usually comma or pipe separated
            return content

        elif 'experience' in section_lower or 'project' in section_lower:
            # Use bullet point format with action verb
            return ContentFormatter.format_bullet_point(content)

        elif 'summary' in section_lower:
            # Summary should flow as prose
            return content

        else:
            return content
